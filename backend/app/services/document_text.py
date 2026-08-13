"""Best-effort text extraction for uploaded CVs and forms.

Untrusted input: uploaded files are never trusted as code or executed;
extraction is read-only and defensively bounded (page/byte caps) so a
malformed or huge file degrades to an empty/partial string instead of
hanging the request.

Deterministic extraction (PyMuPDF, python-docx, openpyxl) covers every
format that actually HAS a text layer. It cannot read pixels: a phone photo
of a printed CV, or a PDF that's a scanned image with no embedded text (very
common for "old records"), yields nothing no matter which library is used.
``extract_text_with_ocr_fallback`` is the one place that gap is closed, via
``LLMProvider.transcribe_image`` -- deterministic extraction always runs
first and is used as-is whenever it finds real text; a vision model is only
invoked for the pages/files that are genuinely just images.
"""

from __future__ import annotations

import io
import logging

from .llm import LLMProvider

logger = logging.getLogger(__name__)

MAX_PDF_PAGES = 40
MAX_TEXT_CHARS = 60_000

# A page/file this short after deterministic extraction is treated as "no
# real text layer" and routed to OCR instead -- PDF export artifacts (page
# numbers, a stray header) can leave a few characters behind even on a fully
# scanned page, so this isn't held to a strict zero.
OCR_FALLBACK_THRESHOLD_CHARS = 20
MAX_OCR_PAGES = 8


def extract_text(content: bytes, mime_type: str) -> str:
    try:
        if mime_type == "application/pdf":
            return _extract_pdf(content)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return _extract_docx(content)
        if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return _extract_xlsx(content)
        if mime_type.startswith("text/"):
            return content.decode("utf-8", errors="ignore")[:MAX_TEXT_CHARS]
    except Exception as exc:  # noqa: BLE001 - defensive: never let a bad upload crash the request
        logger.warning("document_text_extraction_failed", extra={"mime_type": mime_type, "error": str(exc)})
    return ""


async def extract_text_with_ocr_fallback(content: bytes, mime_type: str, llm: LLMProvider) -> str:
    """Same contract as ``extract_text``, plus a vision-model OCR path for
    image uploads and image-only PDF pages. Falls back to the empty string
    (same as ``extract_text``) when no provider is configured -- callers
    already treat an empty result as "nothing readable found" and fail the
    job with a clear message rather than fabricating activities.
    """

    if mime_type in {"image/jpeg", "image/png"}:
        text = await llm.transcribe_image(
            image_bytes=content, mime_type=mime_type,
            instruction="Transcribe every word of readable text in this image exactly as written, preserving line breaks. Output only the transcription, nothing else.",
        )
        return (text or "")[:MAX_TEXT_CHARS]

    if mime_type != "application/pdf":
        return extract_text(content, mime_type)

    try:
        return await _extract_pdf_with_ocr_fallback(content, llm)
    except Exception as exc:  # noqa: BLE001 - defensive: never let a bad upload crash the request
        logger.warning("document_text_extraction_failed", extra={"mime_type": mime_type, "error": str(exc)})
        return ""


async def _extract_pdf_with_ocr_fallback(content: bytes, llm: LLMProvider) -> str:
    import pymupdf

    parts: list[str] = []
    with pymupdf.open(stream=content, filetype="pdf") as document:
        for index, page in enumerate(document):
            if index >= MAX_PDF_PAGES:
                break
            page_text = page.get_text()
            if len(page_text.strip()) >= OCR_FALLBACK_THRESHOLD_CHARS or index >= MAX_OCR_PAGES:
                # Either this page has a real text layer, or it's a scanned
                # page past the OCR budget -- skipped rather than billed
                # per-page against a provider for a huge scanned document.
                parts.append(page_text)
                continue
            # 150 DPI is plenty for OCR legibility and meaningfully cheaper in
            # vision-model tokens than 200+ DPI would be -- worth keeping low
            # given how tight the free-tier tokens-per-minute budget is.
            pixmap = page.get_pixmap(dpi=150)
            transcribed = await llm.transcribe_image(
                image_bytes=pixmap.tobytes("png"), mime_type="image/png",
                instruction="Transcribe every word of readable text on this scanned document page exactly as written, preserving line breaks. Output only the transcription, nothing else.",
            )
            parts.append(transcribed or page_text)
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _extract_pdf(content: bytes) -> str:
    import pymupdf

    parts: list[str] = []
    with pymupdf.open(stream=content, filetype="pdf") as document:
        for index, page in enumerate(document):
            if index >= MAX_PDF_PAGES:
                break
            parts.append(page.get_text())
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _extract_xlsx(content: bytes) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        max_row = min(sheet.max_row or 0, 2000)
        max_col = min(sheet.max_column or 0, 100)
        for row in range(1, max_row + 1):
            cells = [sheet.cell(row=row, column=col).value for col in range(1, max_col + 1)]
            texts = [str(value).strip() for value in cells if value is not None and str(value).strip()]
            if texts:
                parts.append(" | ".join(texts))
    return "\n".join(parts)[:MAX_TEXT_CHARS]
